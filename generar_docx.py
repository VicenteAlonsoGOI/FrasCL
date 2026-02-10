from docx import Document
from docx.shared import Pt
import os

def create_manual():
    doc = Document()
    doc.add_heading('Manual de Usuario: Automatización de Facturas (FrasCL)', 0)
    
    doc.add_heading('Requisitos Previos', level=1)
    p = doc.add_paragraph()
    p.add_run('1. Archivo de Datos: ').bold = True
    p.add_run('Debes tener el archivo Excel con las facturas en la misma carpeta que el programa. El archivo debe llamarse exactamente: 20260209 BORRADOR FRASCL.xlsx.')
    
    p = doc.add_paragraph()
    p.add_run('2. Preparación del Excel: ').bold = True
    p.add_run('Asegúrate de que el Excel no esté abierto en ese momento para evitar errores de acceso.')
    
    doc.add_heading('Cómo ejecutar el programa', level=1)
    doc.add_paragraph('1. Abrir la carpeta compartida: Navega en tu explorador de archivos hasta: \\\\srv-gesico\\Proyectos IA Vicente\\PROYECTO FrasCL')
    doc.add_paragraph('2. Ejecutar el proceso: Busca el archivo llamado INICIAR_AUTOMATIZACION.bat. Haz doble clic sobre él. Se abrirá una ventana negra que te informará del progreso.')
    doc.add_paragraph('3. Resultado: El programa creará automáticamente una carpeta llamada Facturas_Generadas. Dentro encontrarás un archivo PDF por cada cliente.')
    
    doc.add_heading('Notas Importantes', level=1)
    doc.add_paragraph('- Formato de los datos: Si el Excel tiene errores, el PDF mostrará exactamente lo que ponga el Excel.')
    doc.add_paragraph('- Nombres de Clientes: Caracteres especiales como / se sustituirán por _ .')
    doc.add_paragraph('- Columna Total Factura: Se ajusta automáticamente para evitar cortes.')
    
    doc.save('MANUAL_USUARIO.docx')

def create_tech_doc():
    doc = Document()
    doc.add_heading('Documentación Técnica: Sistema de Automatización FrasCL', 0)
    
    doc.add_heading('1. Arquitectura del Sistema', level=1)
    doc.add_paragraph('El sistema sigue un flujo de procesamiento por lotes lineal:')
    doc.add_paragraph('1. Entrada: Lectura de archivo Excel (.xlsx) mediante pandas.', style='List Number')
    doc.add_paragraph('2. Transformación: Limpieza de datos (normalización de moneda, fechas y tipos numéricos).', style='List Number')
    doc.add_paragraph('3. Agrupación: Segmentación de datos por el campo CLIENTE.', style='List Number')
    doc.add_paragraph('4. Generación: Creación de documentos PDF mediante el motor ReportLab.', style='List Number')
    doc.add_paragraph('5. Persistencia: Almacenamiento organizado en carpetas locales y/o de red.', style='List Number')
    
    doc.add_heading('2. Detalles de Implementación (generar_facturas.py)', level=1)
    doc.add_heading('2.1 Procesamiento de Moneda (clean_currency)', level=2)
    doc.add_paragraph('Se implementó una lógica de normalización de strings para manejar importes con € y formatos europeos:')
    doc.add_paragraph('- Elimina símbolos y espacios.')
    doc.add_paragraph('- Normaliza separadores de miles y decimales.')
    doc.add_paragraph('- Gestiona valores nulos devolviendo 0.0.')
    
    doc.add_heading('2.2 Gestión de Fechas', level=2)
    doc.add_paragraph('Resuelve inconsistencias de formato (puntos vs barras) estandarizando a DD/MM/YYYY.')
    
    doc.add_heading('2.3 Generación de PDF y Layout', level=2)
    doc.add_paragraph('- Orientación Landscape (A4).')
    doc.add_paragraph('- Word Wrap: Salto de línea automático para textos largos.')
    doc.add_paragraph('- Anchos Fijos: Garantizan integridad visual en la impresión.')
    
    doc.add_heading('3. Lanzador Automatizado (.bat)', level=1)
    doc.add_paragraph('Incluye soporte para rutas UNC (red) y auto-instalación de librerías.')
    
    doc.add_heading('4. Mantenimiento', level=1)
    doc.add_paragraph('Para actualizar el programa con nuevos datos, simplemente reemplace el archivo Excel manteniendo las palabras clave en los encabezados.')
    
    doc.save('DOCUMENTACION_TECNICA.docx')

if __name__ == "__main__":
    create_manual()
    create_tech_doc()
    print("Documentos Word generados con éxito.")
