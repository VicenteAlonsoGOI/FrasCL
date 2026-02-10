from docx import Document
import os

def update_manual():
    doc_path = 'MANUAL_USUARIO.docx'
    if not os.path.exists(doc_path):
        print("Error: No se encuentra MANUAL_USUARIO.docx")
        return

    doc = Document(doc_path)
    
    replacements = {
        '20260209 BORRADOR FRASCL.xlsx': 'FRASCL.xlsx',
        '\\\\srv-gesico\\Proyectos IA Vicente\\PROYECTO FrasCL': 'la carpeta donde tengas instalado el programa'
    }

    print("Iniciando actualización del manual...")
    
    # Recorrer parrafos
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                print(f"Reemplazando en párrafo: {old_text} -> {new_text}")
                para.text = para.text.replace(old_text, new_text)

    # Recorrer tablas si las hay (a veces se usa para layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                             print(f"Reemplazando en tabla: {old_text} -> {new_text}")
                             para.text = para.text.replace(old_text, new_text)

    doc.save('MANUAL_USUARIO.docx')
    print("Manual actualizado correctamente conservando las imágenes.")

def update_tech_doc():
    # Para la documentación técnica, si no tiene imágenes, podemos regenerarla o actualizarla.
    # Dado que generar_docx.py la crea desde cero, podemos usar ese o actualizarla aqui.
    # Vamos a regenerarla con generar_docx.py para asegurar limpieza, ya que el usuario no dijo que esa tuviera capturas.
    pass

if __name__ == "__main__":
    update_manual()
